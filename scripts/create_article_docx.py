import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
from datetime import datetime

# 设置中文字体
plt.rcParams['font.sans-serif'] = ['Arial Unicode MS', 'SimHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

# 1. 生成封面图
def create_cover():
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis('off')
    
    # 背景
    ax.add_patch(mpatches.Rectangle((0, 0), 10, 6, facecolor='#1a1a2e', edgecolor='none'))
    
    # 标题
    ax.text(5, 5, 'AI_OS', fontsize=40, ha='center', va='center', color='#4cc9f0', weight='bold')
    ax.text(5, 4.2, '多Agent协作系统', fontsize=24, ha='center', va='center', color='#f72585')
    
    # 四个Agent圆形
    agents = [('CEO', 2, 2.5), ('CFO', 4, 2.5), ('CTO', 6, 2.5), ('CMO', 8, 2.5)]
    colors = ['#4361ee', '#3f37c9', '#7209b7', '#560bad']
    
    for (name, x, y), color in zip(agents, colors):
        circle = plt.Circle((x, y), 0.5, facecolor=color, edgecolor='white', linewidth=2)
        ax.add_patch(circle)
        ax.text(x, y, name, fontsize=12, ha='center', va='center', color='white', weight='bold')
    
    # 连接线
    for i in range(len(agents)-1):
        ax.plot([agents[i][1]+0.5, agents[i+1][1]-0.5], [agents[i][2], agents[i+1][2]], 
               color='#4cc9f0', linewidth=2, linestyle='--')
    
    # 作者
    ax.text(5, 0.5, 'AI创世者', fontsize=14, ha='center', va='center', color='#b5179e')
    
    plt.tight_layout()
    
    # 保存
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#1a1a2e')
    buf.seek(0)
    plt.close()
    
    return buf

# 2. 生成三层架构图
def create_architecture():
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.set_xlim(0, 8)
    ax.set_ylim(0, 5)
    ax.axis('off')
    
    # 三层盒子
    layers = [
        ('决策层\nCEO/CFO/CTO/CMO\n（讨论型）', 1, 3.5, '#4361ee'),
        ('管理层\nPM\n（拆解任务）', 1, 2, '#7209b7'),
        ('执行层\nDeveloper/Designer/QA/Operator\n（有工具）', 1, 0.5, '#f72585')
    ]
    
    for (text, x, y, color) in layers:
        rect = mpatches.FancyBboxPatch((x, y), 6, 1, boxstyle="round,pad=0.1", 
                                        facecolor=color, edgecolor='white', linewidth=2, alpha=0.8)
        ax.add_patch(rect)
        ax.text(x+3, y+0.5, text, fontsize=11, ha='center', va='center', color='white', weight='bold')
    
    # 箭头
    ax.annotate('', xy=(3.5, 3.4), xytext=(3.5, 2.6), arrowprops=dict(arrowstyle='->', color='white', lw=2))
    ax.annotate('', xy=(3.5, 1.9), xytext=(3.5, 1.1), arrowprops=dict(arrowstyle='->', color='white', lw=2))
    
    ax.set_title('三层架构', fontsize=16, color='#1a1a2e', pad=10)
    
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    
    return buf

# 3. 生成发言队列对比图
def create_queue_diagram():
    fig, axes = plt.subplots(1, 2, figsize=(8, 3))
    
    # 错误做法
    ax1 = axes[0]
    ax1.set_xlim(0, 4)
    ax1.set_ylim(0, 3)
    ax1.axis('off')
    ax1.set_title('错误：替换队列', fontsize=12, color='red')
    
    # 原队列
    ax1.add_patch(mpatches.Rectangle((0.5, 2), 3, 0.5, facecolor='#ddd', edgecolor='black'))
    ax1.text(2, 2.25, 'CEO: [CTO, CFO, CMO]', fontsize=9, ha='center')
    
    # 被覆盖
    ax1.add_patch(mpatches.Rectangle((0.5, 1), 3, 0.5, facecolor='#fbb', edgecolor='red'))
    ax1.text(2, 1.25, 'CTO: [CEO] CMO没了!', fontsize=9, ha='center', color='red')
    
    # 正确做法
    ax2 = axes[1]
    ax2.set_xlim(0, 4)
    ax2.set_ylim(0, 3)
    ax2.axis('off')
    ax2.set_title('正确：追加队列', fontsize=12, color='green')
    
    # 原队列
    ax2.add_patch(mpatches.Rectangle((0.5, 2), 3, 0.5, facecolor='#ddd', edgecolor='black'))
    ax2.text(2, 2.25, 'CEO: [CTO, CFO, CMO]', fontsize=9, ha='center')
    
    # 追加后
    ax2.add_patch(mpatches.Rectangle((0.5, 1), 3, 0.5, facecolor='#bfb', edgecolor='green'))
    ax2.text(2, 1.25, 'CTO发言后: [CFO, CMO]', fontsize=9, ha='center', color='green')
    
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    
    return buf

# 4. 创建Word文档
def create_word_doc(cover_buf, arch_buf, queue_buf, output_path):
    doc = Document()
    
    # 封面
    doc.add_picture(cover_buf, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 标题
    title = doc.add_heading('我搭了一个四Agent协作系统', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('多Agent讨论是怎么跑起来的？')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    
    # 正文内容
    doc.add_heading('为什么做多Agent？', level=1)
    doc.add_paragraph('单Agent只能"回答问题"，多Agent才能"讨论决策"。企业里很多决策不是一个人说了算——产品要过技术可行性评估，预算要CFO审批，推广要CMO确认。单Agent模拟不了这个过程。所以我搭了AI_OS：四个Agent（CEO/CFO/CTO/CMO）围绕一个议题讨论，动态发言、达成共识、CEO总结。')
    
    doc.add_heading('三层架构', level=1)
    doc.add_picture(arch_buf, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('决策层只讨论，不干活。执行层有工具（terminal、write_file等），能真正执行任务。')
    
    doc.add_heading('动态发言队列：最大的坑', level=1)
    doc.add_paragraph('一开始我让Agent返回next_speakers指定下一位发言者，结果CMO永远轮不到——因为队列被覆盖了。')
    
    doc.add_picture(queue_buf, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('共识检测：Agent总说"不满意"', level=1)
    doc.add_paragraph('Agent返回的satisfied字段永远false，导致讨论不结束。解决方法：双重判断。1. satisfied计数：>=3个Agent满意就算达成。2. 关键词检测：发言中出现"同意/赞同/一致"也算���两个条件满足任意一个就触发共识。')
    
    doc.add_heading('这套架构适合什么场景？', level=1)
    items = [
        '多角色讨论决策（投资委员会、产品评审）',
        '复杂任务拆解执行（需求→设计→开发→测试→上线）',
        '业务流程自动化（收集→汇总→审批→分发）'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph('不适合：简单问答、单线程任务。')
    
    doc.add_heading('后续', level=1)
    doc.add_paragraph('我正在给执行层加真实工具能力。等完成后，这就是一个能讨论、能干活的AI团队。')
    
    # 作者署名
    doc.add_paragraph()
    author = doc.add_paragraph('作者：AI创世者，企业AI化顾问。帮企业把重复业务流程AI化，从诊断到上线全程负责。')
    author.runs[0].font.italic = True
    
    website = doc.add_paragraph('体验AI_OS多Agent讨论：aios.renmin.ai（即将开放）')
    website.runs[0].font.italic = True
    
    # 保存
    doc.save(output_path)
    return output_path

# 执行
if __name__ == '__main__':
    output_dir = '/Users/agent/Documents/任民/文章'
    timestamp = datetime.now().strftime('%Y-%m-%d %H%M')
    output_path = os.path.join(output_dir, f'{timestamp}-我���了一个四Agent协作系统-带图版.docx')
    
    print("生成封面...")
    cover = create_cover()
    
    print("生成架构图...")
    arch = create_architecture()
    
    print("生成队列对比图...")
    queue = create_queue_diagram()
    
    print("组装Word文档...")
    result = create_word_doc(cover, arch, queue, output_path)
    
    print(f"\n完成: {result}")